"""
generate_release_guide.py
Generates Securiti_Desktop_Release_Guide.docx — step-by-step deployment reference
for publishing a new Securiti Desktop Server version.
Run: python generate_release_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours (hex strings for XML, RGBColor for font/run colouring) ──────
NAVY        = RGBColor(0x06, 0x00, 0x61)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_NAVY  = RGBColor(0xE8, 0xE8, 0xF5)
MID_GREY    = RGBColor(0x60, 0x60, 0x60)
GREEN       = RGBColor(0x16, 0x80, 0x3A)
AMBER       = RGBColor(0xD9, 0x77, 0x06)
RED         = RGBColor(0xDC, 0x26, 0x26)
CODE_BG     = RGBColor(0xF3, 0xF4, 0xF6)
CODE_FG     = RGBColor(0x11, 0x18, 0x27)

# Hex strings used for XML cell shading (set_cell_bg)
HEX_NAVY       = '060061'
HEX_WHITE      = 'FFFFFF'
HEX_LIGHT_NAVY = 'E8E8F5'
HEX_CODE_BG    = 'F3F4F6'
HEX_AMBER_BG   = 'FFF7E6'

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, border_type='all', color='BFBFBF', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = ['top', 'left', 'bottom', 'right'] if border_type == 'all' else [border_type]
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_table_borders(table, color='BFBFBF', sz=4):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color, sz=sz)

def navy_band(doc, text):
    """Full-width navy section header band."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, HEX_NAVY)
    set_cell_border(cell, color='060061', sz=6)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text.upper())
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()

def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(13 if level == 2 else 11)
    run.font.color.rgb = NAVY
    run.font.name  = 'Calibri'
    return p

def add_body(doc, text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    run.font.name  = 'Calibri'
    return p

def add_note(doc, text, colour=HEX_LIGHT_NAVY, border_colour='060061'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, colour)
    set_cell_border(cell, color=border_colour, sz=4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    doc.add_paragraph()

def add_warning(doc, text):
    add_note(doc, f'⚠  {text}', colour=HEX_AMBER_BG, border_colour='D97706')

def add_code_block(doc, code: str):
    """Monospaced grey code block."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, HEX_CODE_BG)
    set_cell_border(cell, color='D1D5DB', sz=4)
    cell.paragraphs[0].clear()
    for i, line in enumerate(code.strip().split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_FG
    doc.add_paragraph()

def add_step_header(doc, step_num: int, title: str):
    """Numbered step header with navy left border."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # Number cell
    num_cell = table.cell(0, 0)
    set_cell_bg(num_cell, HEX_NAVY)
    set_cell_border(num_cell, color='060061', sz=6)
    num_cell.width = Cm(1.2)
    num_p = num_cell.paragraphs[0]
    num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_p.paragraph_format.space_before = Pt(4)
    num_p.paragraph_format.space_after  = Pt(4)
    run = num_p.add_run(str(step_num))
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(13)
    run.font.name = 'Calibri'
    # Title cell
    title_cell = table.cell(0, 1)
    set_cell_bg(title_cell, HEX_LIGHT_NAVY)
    set_cell_border(title_cell, color='060061', sz=6)
    title_p = title_cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after  = Pt(4)
    run2 = title_p.add_run(title)
    run2.bold = True
    run2.font.color.rgb = NAVY
    run2.font.size = Pt(12)
    run2.font.name = 'Calibri'
    doc.add_paragraph()

def add_checklist_item(doc, text, done=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    prefix_run = p.add_run('☑  ' if done else '☐  ')
    prefix_run.font.color.rgb = GREEN if done else MID_GREY
    prefix_run.font.size = Pt(10)
    prefix_run.font.name = 'Calibri'
    text_run = p.add_run(text)
    text_run.font.size = Pt(10)
    text_run.font.name = 'Calibri'
    if done:
        text_run.font.color.rgb = MID_GREY

def add_two_col_table(doc, rows, header_row=None, col_widths=None):
    n_cols = len(rows[0]) if rows else 2
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'
    if header_row:
        hrow = table.add_row()
        for i, text in enumerate(header_row):
            cell = hrow.cells[i]
            set_cell_bg(cell, HEX_NAVY)
            set_cell_border(cell, color='060061', sz=4)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
    for r_data in rows:
        row = table.add_row()
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            set_cell_bg(cell, HEX_LIGHT_NAVY if i == 0 else HEX_WHITE)
            set_cell_border(cell, color='BFBFBF', sz=4)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def add_three_col_table(doc, rows, header_row=None):
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    if header_row:
        hrow = table.add_row()
        for i, text in enumerate(header_row):
            cell = hrow.cells[i]
            set_cell_bg(cell, HEX_NAVY)
            set_cell_border(cell, color='060061', sz=4)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
    for r_data in rows:
        row = table.add_row()
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            set_cell_bg(cell, HEX_LIGHT_NAVY if i == 0 else HEX_WHITE)
            set_cell_border(cell, color='BFBFBF', sz=4)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            if i == 0:
                run.bold = True
    doc.add_paragraph()


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover ─────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('SECURITI DESKTOP SERVER')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Release & Deployment Guide')
run2.font.size = Pt(14)
run2.font.color.rgb = MID_GREY
run2.font.name = 'Calibri'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Version 1.0  ·  Colours Universal Inc  ·  June 2026')
run3.font.size = Pt(9)
run3.font.color.rgb = MID_GREY
run3.font.name = 'Calibri'
run3.italic = True

doc.add_paragraph()
doc.add_paragraph()

add_note(doc,
    'Use this guide every time you ship a new version of Securiti Desktop Server. '
    'Follow Steps 1–7 in order. The website updates automatically after Step 6 — '
    'no website code changes are required.',
    colour=HEX_LIGHT_NAVY)

doc.add_paragraph()

# ── Section 1: Overview ───────────────────────────────────────────────────────

navy_band(doc, '1  How Versioning Works')

add_body(doc,
    'There is one source of truth for the version number: MyAppVersion in '
    'SecuritiInstallerV1.iss. Every other system reads from it automatically.')

add_two_col_table(doc,
    rows=[
        ('MyAppVersion in .iss',        'The version number you change — e.g. "1.3.0"'),
        ('Installer filename',           'Auto-derives: Securiti_Desktop_Setup_1.3.0.exe'),
        ('Windows Apps & Features',      'Shows "Securiti Desktop 1.3.0" automatically'),
        ('Installer → right-click → Details', 'Shows version 1.3.0 in PE file properties'),
        ('GitHub release tag',           'You set this manually to match: v1.3.0'),
        ('Website download page',        'Reads from GitHub API — updates automatically'),
    ],
    header_row=['Component', 'Where the version comes from'],
)

add_note(doc,
    'Rule: bump MyAppVersion → build → upload to GitHub Releases. '
    'The website needs no edits.')

doc.add_paragraph()

# ── Section 2: Step-by-step ───────────────────────────────────────────────────

navy_band(doc, '2  Deployment Steps')

# Step 1
add_step_header(doc, 1, 'Bump the version number — one change, one file')
add_body(doc, 'File:  python/securiti-desktop/installer_build/SecuritiInstallerV1.iss')
add_code_block(doc, '#define MyAppVersion  "1.3.0"    ; ← change this line only')
add_note(doc,
    'This is the only manual version edit in the entire release process. '
    'Everything else derives from it automatically.')

# Step 2
add_step_header(doc, 2, 'Build the Flask / Python server')
add_body(doc, 'Run from: python/securiti-desktop/')
add_code_block(doc, 'pyinstaller SecuritiFlaskServerV1.spec')
add_body(doc, 'Output:  dist\\SecuritiFlaskServer\\', space_after=8)

# Step 3
add_step_header(doc, 3, 'Build the C++ video server (Release build only)')
add_body(doc, 'Run from: CPP/video_server/')
add_code_block(doc,
    'cmake -B build -S . -DCMAKE_BUILD_TYPE=Release\n'
    'cmake --build build --config Release')
add_body(doc, 'Output:  build\\Release\\video_server.exe  +  all DLL files')
add_warning(doc,
    'Always build Release — never Debug. Debug builds link to MSVCP140D.dll which '
    'is not redistributable and will fail to start on customer machines.')

# Step 4
add_step_header(doc, 4, 'Populate FilesToInstall\\')
add_body(doc,
    'Copy build outputs into the folder structure below before compiling the installer. '
    'All folders must be present or Inno Setup will error.')
add_code_block(doc,
    'installer_build\\FilesToInstall\\\n'
    '├── Server\\       ← dist\\SecuritiFlaskServer\\  (entire folder)\n'
    '├── VideoServer\\  ← build\\Release\\video_server.exe + *.dll + models\\\n'
    '│                    (copy all .dll files — do NOT copy .pdb debug symbols)\n'
    '├── Nginx\\        ← C:\\nginx\\  (nginx.exe + conf\\ + html\\)\n'
    '├── NSSM\\         ← nssm.exe\n'
    '└── Redist\\       ← vc_redist.x64.exe')

# Step 5
add_step_header(doc, 5, 'Compile the installer with Inno Setup')
add_body(doc,
    'Open Inno Setup Compiler → File → Open → '
    'installer_build\\SecuritiInstallerV1.iss → click Compile (or press F9).')
add_body(doc,
    'Output file will be created automatically at:')
add_code_block(doc,
    'installer_build\\Output\\Securiti_Desktop_Setup_{version}.exe\n'
    '\n'
    'Example for v1.3.0:\n'
    'installer_build\\Output\\Securiti_Desktop_Setup_1.3.0.exe')
add_note(doc,
    'The filename is generated automatically from MyAppVersion — '
    'you do not need to rename it.')

# Step 6
add_step_header(doc, 6, 'Upload to GitHub Releases')
add_body(doc,
    'A ready-to-run PowerShell script handles the entire release in one go: '
    'creating the draft, uploading the installer, and publishing. '
    'There is only one line to edit before running it.')
add_body(doc, 'Script file (open this in any text editor to edit it):')
add_code_block(doc,
    'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build\\publish_release.ps1')
add_body(doc, 'The only line you need to edit is at the top of the script:')
add_code_block(doc,
    '$Version = "1.3.0"   # ← change this to match MyAppVersion in SecuritiInstallerV1.iss')
add_body(doc,
    'Optionally update $ReleaseNotes in the same file to describe what changed in this version. '
    'Then run the script from PowerShell:')
add_code_block(doc,
    'cd "C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build"\n'
    '.\\publish_release.ps1')
add_body(doc, 'The script runs three steps and prints progress:')
add_two_col_table(doc,
    rows=[
        ('[1/3] Creating draft release',  'Creates the release on GitHub with the correct tag and release notes'),
        ('[2/3] Uploading installer',      'Uploads Securiti_Desktop_Setup_{version}.exe — allow 5–15 min for ~400 MB'),
        ('[3/3] Publishing release',       'Takes the release out of draft so it is publicly visible'),
    ],
    header_row=['Step', 'What it does'],
)
add_body(doc, 'On success the script prints the release URL and download URL.')
add_warning(doc,
    'Do not close PowerShell during the upload. If the upload fails mid-way, '
    'delete the draft release on GitHub (releases page → Delete) then re-run the script.')

# Step 7
add_step_header(doc, 7, 'Verify the website has updated')
add_body(doc,
    'Open the download page in a browser. Within a few minutes of publishing the '
    'GitHub release the page should show the new version number, release date, '
    'and file size automatically. No website edits are needed.')
add_body(doc, 'To confirm the API is returning the new release, open this URL in a browser:')
add_code_block(doc,
    'https://api.github.com/repos/dominic96/securiti-releases/releases/latest')
add_body(doc, 'Check that these fields match your new build:')
add_two_col_table(doc,
    rows=[
        ('tag_name',          'Should be "v1.3.0" (or your new version)'),
        ('assets[0].name',    'Should be "Securiti_Desktop_Setup_1.3.0.exe"'),
        ('assets[0].size',    'Should match the file size of the new installer (bytes)'),
        ('assets[0].browser_download_url', 'Should include "v1.3.0" in the URL path'),
    ],
    header_row=['API field', 'Expected value'],
)

doc.add_paragraph()

# ── Section 3: How Upgrades Work ─────────────────────────────────────────────

navy_band(doc, '3  How Upgrades Work — What the Installer Does on an Existing Machine')

add_body(doc,
    'When a customer runs a newer installer on a machine that already has Securiti Desktop '
    'installed, Inno Setup detects the existing installation and runs in upgrade mode. '
    'The behaviour is different from a fresh install in every way that matters.')

add_note(doc,
    'No data is lost during an upgrade. The database, cameras, faces, identities, '
    'notifications, stream tokens, and all user accounts survive intact.')

add_heading(doc, 'What the installer does, step by step (upgrade mode)')

add_two_col_table(doc,
    rows=[
        ('1. Detect existing install',
         'Inno Setup checks the Windows registry for the app\'s uninstall key '
         '(keyed to the AppId GUID). If present → upgrade mode. If absent → fresh install.'),
        ('2. Stop all three services',
         'sc.exe stops VideoServerSecuritiService, FlaskSecuritiService, and '
         'NginxSecuritiService in reverse dependency order. This releases the exclusive '
         'file lock each service holds on its own .exe, allowing files to be overwritten.'),
        ('3. Overwrite binary files',
         'All application files are replaced: video_server.exe, SecuritiFlaskServer.exe, '
         'nginx.exe, all DLLs, all AI model files, and Nginx config. Only software files '
         'in Program Files are touched — nothing in ProgramData is written to.'),
        ('4. Skip nssm install',
         'The nssm install commands are skipped (they are guarded with Check: IsFreshInstall). '
         'The services are already registered — re-registering them would error.'),
        ('5. Update service config',
         'All nssm set commands still run. These update service properties (AppDirectory, '
         'log paths, rotation settings) in case anything changed between versions.'),
        ('6. Restart services',
         'nssm start runs for all three services in order: Nginx → Flask → VideoServer. '
         'Services come back up running the new binaries.'),
        ('7. Show upgrade confirmation',
         '"Securiti Desktop updated to v1.x.x successfully. Your cameras, users, '
         'identities, and all data have been preserved."'),
    ],
    header_row=['Phase', 'What happens'],
)

add_heading(doc, 'What is NEVER touched during an upgrade')

add_two_col_table(doc,
    rows=[
        ('securiti.db',
         'The SQLite database at C:\\ProgramData\\Securiti Desktop\\data\\securiti.db. '
         'Contains: cameras, stream tokens, face embeddings, identities, identity sightings, '
         'security zones, security rules, security alerts, notifications, sessions.'),
        ('ServerLogs\\ and VideoServerLogs\\',
         'Log files are never deleted or truncated by the installer.'),
        ('Firebase credentials',
         'Any credential or key files stored in ProgramData are untouched.'),
        ('LocalX tunnel config',
         'Tunnel registration and config files bundled in the Flask server folder '
         'survive because the installer uses ignoreversion — it overwrites the binary '
         'but leaves config files the server wrote at runtime alone.'),
    ],
    header_row=['Data', 'Location / Detail'],
)

add_heading(doc, 'What IS replaced during an upgrade')

add_two_col_table(doc,
    rows=[
        ('video_server.exe',             'C:\\Program Files\\Securiti Desktop\\VideoServer\\'),
        ('SecuritiFlaskServer.exe + all PyInstaller files', 'C:\\Program Files\\Securiti Desktop\\Server\\'),
        ('nginx.exe + nginx.conf',       'C:\\Program Files\\Securiti Desktop\\Nginx\\'),
        ('All DLL files',                'C:\\Program Files\\Securiti Desktop\\VideoServer\\*.dll'),
        ('All AI model files (.onnx)',   'C:\\Program Files\\Securiti Desktop\\VideoServer\\models\\'),
        ('nssm.exe',                     'C:\\Program Files\\Securiti Desktop\\NSSM\\'),
    ],
    header_row=['File(s)', 'Location'],
)

add_heading(doc, 'How the installer tells fresh install from upgrade')

add_body(doc,
    'Inno Setup writes an uninstall registry key the first time the software is installed. '
    'On every subsequent run the installer checks for this key:')
add_code_block(doc,
    'HKLM\\Software\\Microsoft\\Windows\\CurrentVersion\\Uninstall\\\n'
    '{8F3A2D1E-7B4C-4E9F-A6D2-1C8E5F0B3A97}_is1\n'
    '\n'
    'Key present   → IsUpgrade() = True  → upgrade mode\n'
    'Key absent    → IsUpgrade() = False → fresh install mode')

add_note(doc,
    'The AppId GUID (8F3A2D1E-7B4C-4E9F-A6D2-1C8E5F0B3A97) is fixed in '
    'SecuritiInstallerV1.iss and must never change. Changing it would make the '
    'installer treat every upgrade as a fresh install and wipe the database.')

add_heading(doc, 'Fresh install vs upgrade — side by side')

add_three_col_table(doc,
    rows=[
        ('Detects existing install',       'No (first run)',                      'Yes (registry key present)'),
        ('Stops running services',         'Yes (safe no-op if not installed)',   'Yes (required to unlock files)'),
        ('Wipes securiti.db',              'Yes (clears stale dev data)',          'No — database is preserved'),
        ('Runs nssm install',              'Yes (registers new services)',         'No — services already exist'),
        ('Runs nssm set (config update)',  'Yes',                                  'Yes'),
        ('Starts services',               'Yes',                                  'Yes (restarts them)'),
        ('Success message',               '"Installed successfully"',             '"Updated to vX.X.X — data preserved"'),
    ],
    header_row=['Action', 'Fresh Install', 'Upgrade'],
)

doc.add_paragraph()

# ── Section 4: Quick Reference Checklist ─────────────────────────────────────

navy_band(doc, '4  Release Checklist  (print and tick off each time)')

add_body(doc, 'Use this checklist for every release. Tick each item before moving to the next.')
doc.add_paragraph()

steps_check = [
    ('Step 1', 'MyAppVersion updated in SecuritiInstallerV1.iss'),
    ('Step 2', 'Flask server built with PyInstaller — dist\\SecuritiFlaskServer\\ present'),
    ('Step 3', 'C++ video server built in Release mode — video_server.exe present in build\\Release\\'),
    ('Step 4', 'FilesToInstall\\ fully populated — Server\\, VideoServer\\, Nginx\\, NSSM\\, Redist\\ all present'),
    ('Step 5', 'Inno Setup compiled successfully — Securiti_Desktop_Setup_{version}.exe present in Output\\'),
    ('Step 5', 'Installer filename includes the correct version number'),
    ('Step 6', '$Version updated in publish_release.ps1 to match MyAppVersion'),
    ('Step 6', 'publish_release.ps1 ran without errors — all 3 steps completed'),
    ('Step 6', 'Release URL printed by script — confirm tag is v{version} and not a draft'),
    ('Step 7', 'GitHub API returns the new version at /releases/latest'),
    ('Step 7', 'Website download page shows new version, date, and file size'),
    ('Step 7', 'Download button triggers the correct versioned .exe file'),
]
for label, text in steps_check:
    add_checklist_item(doc, f'[{label}]  {text}')

doc.add_paragraph()

# ── Section 4: What updates automatically ────────────────────────────────────

navy_band(doc, '5  What Updates Automatically vs What You Do Manually')

add_two_col_table(doc,
    rows=[
        ('Installer filename',           '✅ Auto — derived from MyAppVersion in .iss'),
        ('Windows Apps & Features name', '✅ Auto — derived from MyAppVersion in .iss'),
        ('Exe right-click → Details',    '✅ Auto — derived from MyAppVersion in .iss'),
        ('Website version badge',        '✅ Auto — reads from GitHub API'),
        ('Website release date',         '✅ Auto — reads from GitHub API'),
        ('Website file size',            '✅ Auto — reads from GitHub API'),
        ('Website download button URL',  '✅ Auto — reads from GitHub API'),
        ('Website version history table','✅ Auto — reads from GitHub API'),
        ('MyAppVersion in .iss',         '✏️  Manual — you change this (one line)'),
        ('$Version in publish_release.ps1', '✏️  Manual — one line at the top of the script'),
        ('$ReleaseNotes in publish_release.ps1', '✏️  Manual — describe what changed in this version'),
    ],
    header_row=['Component', 'How it updates'],
)

# ── Section 5: Common Mistakes ────────────────────────────────────────────────

navy_band(doc, '6  Common Mistakes to Avoid')

mistakes = [
    ('Building in Debug mode',
     'Always use --config Release. Debug builds link to MSVCP140D.dll which is not '
     'redistributable. The installer will fail to run on customer machines.'),
    ('Changing the AppId GUID in SecuritiInstallerV1.iss',
     'The AppId {8F3A2D1E-7B4C-4E9F-A6D2-1C8E5F0B3A97} must never be changed. '
     'It is the key Inno Setup uses to detect an existing installation (IsUpgrade). '
     'If you change it, every upgrade will be treated as a fresh install — the installer '
     'will wipe securiti.db and all customer data: cameras, faces, identities, '
     'notifications, and accounts will be permanently deleted with no warning.'),
    ('Leaving the release as a Draft',
     'Draft releases are not visible to the public and the GitHub API will not '
     'return them at /releases/latest. Always publish before verifying the website.'),
    ('Hardcoding version numbers in website code',
     'The website reads all values dynamically from the GitHub API. Never hardcode '
     'the version, file size, date, or download URL — it will be wrong for the next release.'),
    ('Mismatching MyAppVersion and the GitHub tag',
     'If MyAppVersion = "1.3.0" but the GitHub tag is "v1.4.0", the filename and '
     'the website will show different versions. They must match exactly.'),
    ('Forgetting to populate FilesToInstall\\',
     'Inno Setup will silently omit files or error if the source folders are stale. '
     'Always re-copy fresh build outputs before compiling the installer.'),
    ('Uploading a Debug .pdb file',
     'Do not copy .pdb symbol files into VideoServer\\. They are large (~200 MB each) '
     'and serve no purpose in the installer.'),
]

for title, body in mistakes:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f'✗  {title}')
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.8)
    run2 = p2.add_run(body)
    run2.font.size = Pt(9.5)
    run2.font.name = 'Calibri'
    run2.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

doc.add_paragraph()

# ── Section 6: Key Paths ──────────────────────────────────────────────────────

navy_band(doc, '7  Key File Paths & URLs')

add_two_col_table(doc,
    rows=[
        ('Version definition',    'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build\\SecuritiInstallerV1.iss  →  #define MyAppVersion'),
        ('Release script',        'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build\\publish_release.ps1  →  $Version'),
        ('Flask build output',    'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\dist\\SecuritiFlaskServer\\'),
        ('C++ build output',      'C:\\Users\\domin\\Workspace\\CPP\\video_server\\build\\Release\\'),
        ('Installer input',       'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build\\FilesToInstall\\'),
        ('Installer output',      'C:\\Users\\domin\\Workspace\\python\\securiti-desktop\\installer_build\\Output\\Securiti_Desktop_Setup_{version}.exe'),
        ('GitHub releases repo',  'https://github.com/dominic96/securiti-releases'),
        ('GitHub API (latest)',   'https://api.github.com/repos/dominic96/securiti-releases/releases/latest'),
        ('GitHub API (all)',      'https://api.github.com/repos/dominic96/securiti-releases/releases'),
        ('Download page fallback','https://github.com/dominic96/securiti-releases/releases/latest'),
    ],
    header_row=['Item', 'Path / URL'],
)

# ── Footer ────────────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Securiti Desktop Server — Release Guide  ·  Colours Universal Inc  ·  confidential')
run.font.size = Pt(8)
run.font.color.rgb = MID_GREY
run.font.name = 'Calibri'
run.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────

import os, time
output = 'Securiti_Desktop_Release_Guide.docx'
tmp    = output + '.tmp.docx'
doc.save(tmp)
for _ in range(5):
    try:
        if os.path.exists(output):
            os.replace(tmp, output)
        else:
            os.rename(tmp, output)
        break
    except PermissionError:
        print('File is open — retrying in 3 s (close Word if prompted)...')
        time.sleep(3)
else:
    print(f'Could not replace {output} — saved as {tmp} instead')
    output = tmp
print(f'Generated: {output}')
